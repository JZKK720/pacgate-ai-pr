"""Bundle add-on documents into PDF and Word packages for client delivery."""
from __future__ import annotations

import subprocess
from pathlib import Path

from PyPDF2 import PdfMerger
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path("scope-assets/contracts/preapproval-contract-draft")
OUT = BASE / "out"
OUT.mkdir(parents=True, exist_ok=True)

PYTHON = "c:/Users/cubecloud-io/github-pr/pacgate-ai-pr/.venv/Scripts/python.exe"
MD2PDF = "safe_markdown_to_pdf.py"

# ── Files to bundle ──────────────────────────────────────────────
ADDON_FILES_ZH = [
    ("addon-services-schedule-zh", "可选服务与附加费用清单"),
    ("remote-advisory-service-order-zh", "远程顾问服务订单"),
    ("annual-license-renewal-order-zh", "年度许可续费订单"),
    ("onsite-service-order-zh", "现场服务订单"),
]

ADDON_FILES_EN = [
    ("addon-services-schedule-en", "Optional Services & Add-On Schedule"),
    ("remote-advisory-service-order-en", "Remote Advisory Service Order"),
    ("annual-license-renewal-order-en", "Annual License Renewal Order"),
    ("onsite-service-order-en", "On-Site Service Order"),
]


def ensure_pdf(stem: str) -> Path:
    """Generate PDF from markdown if not already present."""
    pdf = BASE / f"{stem}.pdf"
    md = BASE / f"{stem}.md"
    if not pdf.exists() or md.stat().st_mtime > pdf.stat().st_mtime:
        subprocess.run([PYTHON, MD2PDF, str(md), str(pdf)], check=True)
    return pdf


def merge_pdfs(stems: list[tuple[str, str]], output_name: str) -> Path:
    """Merge multiple PDFs into one bundle."""
    merger = PdfMerger()
    for stem, _title in stems:
        pdf = ensure_pdf(stem)
        merger.append(str(pdf))
    out_path = OUT / output_name
    merger.write(str(out_path))
    merger.close()
    return out_path


def build_word_bundle(stems: list[tuple[str, str]], output_name: str) -> Path:
    """Convert markdown files to a single Word document with title pages."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for i, (stem, title) in enumerate(stems):
        md_path = BASE / f"{stem}.md"
        if not md_path.exists():
            continue

        # Title page per document
        if i > 0:
            doc.add_page_break()

        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Read markdown and add as paragraphs
        text = md_path.read_text(encoding="utf-8")
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("---"):
                doc.add_paragraph("─" * 60)
            elif line.startswith("|"):
                # Simple table rendering
                p = doc.add_paragraph(line)
                p.style = doc.styles["Normal"]
                run = p.runs[0] if p.runs else p.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.runs[0] if p.runs else p.add_run(line[2:])
                run.font.italic = True
                run.font.color.rgb = None  # default
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line, style="List Bullet")
            elif line.startswith("1. ") or line.startswith("2. "):
                doc.add_paragraph(line, style="List Number")
            else:
                doc.add_paragraph(line)

    out_path = OUT / output_name
    doc.save(str(out_path))
    return out_path


# ── Main ──────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating PDF bundles...")
    zh_pdf = merge_pdfs(ADDON_FILES_ZH, "addon-bundle-zh.pdf")
    en_pdf = merge_pdfs(ADDON_FILES_EN, "addon-bundle-en.pdf")
    print(f"  {zh_pdf.name}  ({zh_pdf.stat().st_size:,} bytes)")
    print(f"  {en_pdf.name}  ({en_pdf.stat().st_size:,} bytes)")

    print("Generating Word bundles...")
    zh_docx = build_word_bundle(ADDON_FILES_ZH, "addon-bundle-zh.docx")
    en_docx = build_word_bundle(ADDON_FILES_EN, "addon-bundle-en.docx")
    print(f"  {zh_docx.name}  ({zh_docx.stat().st_size:,} bytes)")
    print(f"  {en_docx.name}  ({en_docx.stat().st_size:,} bytes)")

    print("\nDone. Bundles in scope-assets/contracts/preapproval-contract-draft/out/")
